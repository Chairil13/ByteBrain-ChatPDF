import streamlit as st
from PyPDF2 import PdfReader
from langchain.text_splitter import RecursiveCharacterTextSplitter
import os
from langchain_google_genai import GoogleGenerativeAIEmbeddings
import google.generativeai as genai
from langchain.vectorstores import FAISS
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.chains.question_answering import load_qa_chain
from langchain.prompts import PromptTemplate
from dotenv import load_dotenv
import re
import difflib
from docx import Document
from io import BytesIO

# Mengatur judul dan ikon
st.set_page_config(page_title="ByteBrain - Chat with PDF", page_icon="🤖")

# Style CSS
st.markdown(
    """
    <style>
    .stApp {
        background-color: #f7dcbe;
    }
    section[data-testid="stSidebar"] {
        background-color: #abdbe3;
    }
    span[style*="background-color: yellow"] {
        padding: 2px;
        border-radius: 3px;
    }
    p {
        color: black;
    }
    </style>
    """,
    unsafe_allow_html=True
)

# Memuat token API Gemini
load_dotenv()
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))
print(os.getenv("GOOGLE_API_KEY"))

# Fungsi untuk mengambil teks dari file PDF yang diunggah.
def get_pdf_text(pdf_docs):
    text_with_page = []
    for pdf in pdf_docs:
        pdf_reader = PdfReader(pdf)
        for page_num, page in enumerate(pdf_reader.pages, 1):
            text = page.extract_text()
            text_with_page.append((text, page_num))
    return text_with_page

# Fungsi untuk membagi teks menjadi potongan-potongan kecil agar lebih mudah diolah.
def get_text_chunks(text_with_page):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=10000, chunk_overlap=1000)
    chunks = []
    for text, page_num in text_with_page:
        split_texts = text_splitter.split_text(text)
        chunks.extend([(chunk, page_num) for chunk in split_texts])
    return chunks

# Fungsi untuk membuat dan menyimpan vector store dan direpsentasi ke embeddings.
def get_vector_store(text_chunks):
    embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
    texts = [chunk[0] for chunk in text_chunks]
    metadatas = [{"page": chunk[1]} for chunk in text_chunks]
    vector_store = FAISS.from_texts(texts, embedding=embeddings, metadatas=metadatas)
    vector_store.save_local("faiss_index")

# Fungsi untuk membuat trained data model atau template
def get_conversational_chain():
    prompt_template = """
    Answer the question as detailed as possible from the provided context, make sure to provide all the details, if the answer is not in
    provided context just say, "Jawaban tidak tersedia dari konteks yang diberikan", don't provide the wrong answer\n\n
    Context:\n {context}?\n
    Question: \n{question}\n

    Answer:
    """

    model = ChatGoogleGenerativeAI(model="gemini-1.5-flash", temperature=0.3) # temperature rendah biar akurat
    prompt = PromptTemplate(template=prompt_template, input_variables=["context", "question"])
    chain = load_qa_chain(model, chain_type="stuff", prompt=prompt)
    return chain

# Fungsi untuk menemukan dan menyoroti teks dalam PDF yang cocok dengan jawaban AI.
def find_and_highlight_text(ai_answer, doc_text, page_num):
    ai_sentences = ai_answer.split('.')
    doc_sentences = doc_text.split('.')

    highlighted_text = []
    for ai_sentence in ai_sentences:
        ai_sentence = ai_sentence.strip()
        if ai_sentence:
            matches = difflib.get_close_matches(ai_sentence, doc_sentences, n=1, cutoff=0.7)
            if matches:
                highlighted_text.append(f"Page {page_num}: <span style='background-color: yellow;'>{matches[0]}</span>")

    return highlighted_text

# Fungsi untuk menjawab pertanyaan user berdasarkan konten PDF.
def answer_question(user_question):
    embeddings = GoogleGenerativeAIEmbeddings(model="models/text-embedding-004")
    
    new_db = FAISS.load_local("faiss_index", embeddings, allow_dangerous_deserialization=True)
    docs = new_db.similarity_search(user_question)

    chain = get_conversational_chain()
    
    response = chain(
        {"input_documents": docs, "question": user_question},
        return_only_outputs=True
    )

    ai_answer = response['output_text']
    st.markdown(f"**Reply:**\n\n{ai_answer}")

    st.markdown("**Jawaban yang relevan dari PDF:**")
    highlighted_texts = []
    for doc in docs:
        highlighted = find_and_highlight_text(ai_answer, doc.page_content, doc.metadata['page'])
        highlighted_texts.extend(highlighted)

    for text in highlighted_texts[:3]:
        st.markdown(text, unsafe_allow_html=True)

    if not highlighted_texts:
        st.markdown("Tidak ditemukan kecocokan jawaban dari PDF.")

# Fungsi untuk memformat kuis menjadi format yang sesuai
def format_quiz(raw_quiz):
    formatted_quiz = []
    questions = re.split(r'\d+\.', raw_quiz)[1:]

    for i, question in enumerate(questions, 1):
        lines = question.strip().split('\n')
        question_text = lines[0].strip()
        
        if question_text.startswith('[Generated Question'):
            continue
        
        options = []
        answer = ""

        for line in lines[1:]:
            line = line.strip()
            if line.startswith(('A.', 'B.', 'C.', 'D.')):
                options.append(line)
            elif 'Answer:' in line:
                answer = line

        formatted_question = f"{i}. Question: {question_text}\n"
        for option in options:
            formatted_question += f"   {option}\n"
        formatted_question += f"   {answer}\n"

        formatted_quiz.append(formatted_question)

    return "\n".join(formatted_quiz)

# Fungsi untuk menghasilkan kuis berdasarkan konteks yang diekstraksi dari PDF
def generate_quiz(context_language, quiz_language, num_questions, difficulty_level):
    difficulty_prompt = {
        "Mudah": "Create simple and straightforward questions.",
        "Sedang": "Create moderately complex questions.",
        "Sulit": "Create challenging and difficult questions."
    }

    question_format = """
    {i}. [Write a specific, relevant question here based on the context]
    A. [Option A]
    B. [Option B]
    C. [Option C]
    D. [Option D]
    Answer: [Correct Answer]
    """
    questions_prompt = "\n\n".join([question_format.format(i=i+1) for i in range(num_questions)])
    
    prompt_template = f"""
    Create a multiple-choice quiz in {quiz_language} based on the given context in {context_language}. Follow these rules strictly:
    1. Ensure each question is relevant and specific to the context provided.
    2. Provide 4 answer options (A, B, C, D) for each question.
    3. Indicate the correct answer after the options.
    4. Do not use placeholders like "[Write a specific, relevant question here based on the context]". Instead, write actual, context-specific questions.
    5. {difficulty_prompt[difficulty_level]}
    6. The output format should be exactly as shown below, with each question on a new line and options labeled A, B, C, D:

    Context:\n {{context}}\n

    Quiz:
    {questions_prompt}
    """
    
    model = ChatGoogleGenerativeAI(model="gemini-1.5-flash", temperature=0.7)
    prompt = PromptTemplate(template=prompt_template, input_variables=["context"])
    chain = load_qa_chain(model, chain_type="stuff", prompt=prompt)
    
    embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
    new_db = FAISS.load_local("faiss_index", embeddings, allow_dangerous_deserialization=True)
    docs = new_db.similarity_search("", k=5)  # Get top 5 most relevant chunks
    context = " ".join([doc.page_content for doc in docs])
    
    response = chain(
        {"input_documents": docs, "question": prompt.format(context=context)},
        return_only_outputs=True
    )
    
    formatted_quiz = format_quiz(response["output_text"])
    
    st.write("Generated Quiz:")
    st.code(formatted_quiz)
    
    # buat dokumen word
    doc = Document()
    doc.add_heading("Generated by ByteBrainAI - Chairil & Samrah", 0)  # Add title
    for line in formatted_quiz.split('\n'):
        doc.add_paragraph(line)
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    st.download_button(
        label="Download Quiz (Word)",
        data=buffer,
        file_name="Generated_Quiz_by_ByteBrain.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# Fungsi utama yang mengatur alur program, mulai dari pengunggahan PDF, ekstraksi teks, hingga pembuatan kuis
def main():  
    st.header("ByteBrain - Chat with PDF 📂")
    st.markdown("Fitur utama dari sistem ini yaitu user dapat memperoleh informasi berdasarkan file PDF yang di upload dan dapat membuat daftar pertanyaan atau kuis beserta jawabannya. Silahkan upload file PDF Anda, form upload terletak pada pojok kiri atas Browser ↖️", unsafe_allow_html=True)

    quiz_language = st.selectbox("Pilih Bahasa Kuis:", ["Indonesian", "English"])
    num_questions = st.number_input("Masukkan jumlah soal yang diinginkan:", min_value=1, max_value=100, value=10)
    difficulty_level = st.selectbox("Pilih Tingkat Kesulitan:", ["Mudah", "Sedang", "Sulit"])

    pdf_uploaded = bool(st.session_state.get("pdf_uploaded", False))
    text_extracted = bool(st.session_state.get("text_extracted", False))

    if st.button("Generate Quiz"):
        if not pdf_uploaded:
            st.warning("Tolong upload PDF terlebih dahulu")
        elif not text_extracted:
            st.warning("Tolong ekstrak teks dari PDF terlebih dahulu")
        else:
            with st.spinner("Generating quiz..."):
                generate_quiz("Indonesian", quiz_language, num_questions, difficulty_level)

    with st.sidebar:
        st.title("Menu:")
        pdf_docs = st.file_uploader("Upload file PDF anda lalu klik tombol Ekstrak Teks dari PDF, setelah itu Anda bisa tanyakan pertanyaan atau generate quiz.", accept_multiple_files=True)

        if pdf_docs:
            is_pdf = all([pdf.type == "application/pdf" for pdf in pdf_docs])
            if is_pdf:
                st.session_state["pdf_uploaded"] = True
            else:
                st.warning("File dokumen tidak didukung, silakan upload file PDF.")
                st.session_state["pdf_uploaded"] = False
        else:
            st.session_state["pdf_uploaded"] = False

        if st.button("Ekstrak Teks"):
            if not pdf_docs:
                st.warning("Tolong upload PDF terlebih dahulu")
            elif not st.session_state.get("pdf_uploaded", False):
                st.warning("File dokumen tidak didukung, silakan upload file PDF.")
            else:
                with st.spinner("Memuat..."):
                    text_with_page = get_pdf_text(pdf_docs)
                    text_chunks = get_text_chunks(text_with_page)
                    get_vector_store(text_chunks)
                    st.session_state["text_extracted"] = True
                    st.success("Selesai")
        
        st.header("Ajukan Pertanyaan:")
        user_question = st.text_input("Masukkan pertanyaan Anda di sini:")
        if st.button("Kirim"):
            if not pdf_uploaded:
                st.warning("Tolong upload PDF terlebih dahulu")
            elif not text_extracted:
                st.warning("Tolong ekstrak teks dari PDF terlebih dahulu")
            else:
                if user_question:
                    if os.path.exists("faiss_index"):
                        with st.spinner("Mencari jawaban..."):
                            answer_question(user_question)
                    else:
                        st.warning("Tolong ekstrak teks dari PDF terlebih dahulu")

# Memastikan fungsi main() dijalankan ketika script di eksekusi
if __name__ == "__main__":
    main()